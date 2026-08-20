from __future__ import annotations

from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Cm


WORKSPACE = Path(r"C:\Users\playdata2\SKN30_FINAL_5")
SOURCE = WORKSPACE / ".docx_work" / "template.docx"
OUTPUT = (
    WORKSPACE
    / "outputs"
    / "[데이터 수집 및 저장] 데이터베이스_저장소 설계 문서_30기_5팀.docx"
)
ERD_PATH = WORKSPACE / ".docx_work" / "decision_erd.png"


def set_paragraph_text(paragraph: Paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def set_cell_text(cell, text: str) -> None:
    paragraph = cell.paragraphs[0]
    set_paragraph_text(paragraph, text)
    for extra in cell.paragraphs[1:]:
        set_paragraph_text(extra, "")


def replace_table_rows(table, rows: list[list[str]]) -> None:
    body_sample = deepcopy(table.rows[1]._tr)
    while len(table.rows) > 1:
        table._tbl.remove(table.rows[-1]._tr)
    for values in rows:
        table._tbl.append(deepcopy(body_sample))
        row = table.rows[-1]
        tr_pr = row._tr.get_or_add_trPr()
        if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
            tr_pr.append(OxmlElement("w:cantSplit"))
        for index, value in enumerate(values):
            set_cell_text(row.cells[index], value)
    header = table.rows[0]
    tr_pr = header._tr.get_or_add_trPr()
    if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblHeader") is None:
        tr_pr.append(OxmlElement("w:tblHeader"))
    if tr_pr.find("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}cantSplit") is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def insert_paragraph_after(paragraph: Paragraph, text: str = "") -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    result = Paragraph(new_p, paragraph._parent)
    if text:
        result.add_run(text)
    return result


def draw_erd(path: Path) -> None:
    canvas = Image.new("RGB", (1800, 1080), "white")
    draw = ImageDraw.Draw(canvas)
    font_regular = ImageFont.truetype(r"C:\Windows\Fonts\malgun.ttf", 34)
    font_small = ImageFont.truetype(r"C:\Windows\Fonts\malgun.ttf", 28)
    font_bold = ImageFont.truetype(r"C:\Windows\Fonts\malgunbd.ttf", 38)
    blue = "#1767c9"
    pale = "#d9eaf8"
    dark = "#1f2937"

    draw.text((70, 35), "Wave 6 기준 핵심 ERD", fill=dark, font=font_bold)

    boxes = {
        "users": (70, 160, 390, 280),
        "routines": (70, 430, 390, 550),
        "daily_contexts": (70, 700, 390, 820),
        "catalog_versions": (520, 160, 880, 280),
        "exercises": (520, 430, 880, 550),
        "decision_runs": (1030, 350, 1410, 490),
        "agent_proposals": (1480, 90, 1770, 210),
        "plan_candidates": (1480, 270, 1770, 390),
        "plan_items": (1480, 450, 1770, 570),
        "safety_reviews": (1480, 630, 1770, 750),
        "decision_options": (1480, 810, 1770, 930),
    }

    for name, (x1, y1, x2, y2) in boxes.items():
        fill = pale if name == "decision_runs" else "#f8fbff"
        draw.rounded_rectangle((x1, y1, x2, y2), radius=16, fill=fill, outline=blue, width=5)
        bbox = draw.textbbox((0, 0), name, font=font_regular)
        tx = (x1 + x2 - (bbox[2] - bbox[0])) / 2
        ty = (y1 + y2 - (bbox[3] - bbox[1])) / 2 - 5
        draw.text((tx, ty), name, fill=dark, font=font_regular)

    def arrow(start, end, label: str = "") -> None:
        draw.line((start, end), fill=blue, width=5)
        ex, ey = end
        sx, sy = start
        if abs(ex - sx) >= abs(ey - sy):
            direction = 1 if ex > sx else -1
            points = [(ex, ey), (ex - 18 * direction, ey - 12), (ex - 18 * direction, ey + 12)]
        else:
            direction = 1 if ey > sy else -1
            points = [(ex, ey), (ex - 12, ey - 18 * direction), (ex + 12, ey - 18 * direction)]
        draw.polygon(points, fill=blue)
        if label:
            mx, my = (sx + ex) / 2, (sy + ey) / 2
            draw.text((mx + 8, my - 38), label, fill=dark, font=font_small)

    arrow((390, 220), (1030, 390), "1:N")
    arrow((390, 490), (1030, 420), "N:1")
    arrow((390, 760), (1030, 450), "N:1")
    arrow((880, 220), (1030, 380), "N:1")
    arrow((700, 280), (700, 430), "1:N")
    arrow((1410, 390), (1480, 150), "1:4")
    arrow((1410, 410), (1480, 330), "1:N")
    arrow((1625, 390), (1625, 450), "1:N")
    arrow((1410, 440), (1480, 690), "1:1")
    arrow((1410, 460), (1480, 870), "1:N")
    arrow((880, 490), (1480, 510), "N:1")

    draw.text(
        (70, 980),
        "결정 입력 snapshot·버전·4개 proposal·후보·Safety veto·최종 option을 분리 저장",
        fill=dark,
        font=font_small,
    )
    canvas.save(path, quality=95)


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    draw_erd(ERD_PATH)
    doc = Document(SOURCE)

    set_cell_text(doc.tables[0].cell(0, 0), "SK 네트웍스 Family AI 30기 : 5팀\n데이터 베이스/저장소 설계 문서")
    cover = doc.tables[1]
    set_cell_text(cover.cell(1, 1), "2026. 8. 14.")
    set_cell_text(
        cover.cell(2, 1),
        "https://github.com/SKNETWORKS-FAMILY-AICAMP/SKN30-FINAL-5Team",
    )
    set_cell_text(cover.cell(3, 1), "채동현")

    by_text = {p.text.strip(): p for p in doc.paragraphs if p.text.strip()}
    set_paragraph_text(
        by_text["모델링 방법 :"],
        "모델링 방법 : 관계형 모델링, SQLAlchemy 2.0 선언형 모델, Alembic 버전 migration",
    )
    set_paragraph_text(
        by_text["정규화 수준 :"],
        "정규화 수준 : 3NF 중심. 재현용 snapshot·proposal·버전 metadata만 JSONB 사용",
    )
    set_paragraph_text(
        by_text["도구 :"],
        "도구 : PostgreSQL, SQLAlchemy 2.0, Alembic, FastAPI/Pydantic, pytest",
    )

    entity_rows = [
        ["사용자", "users", "id", "UUID PK", "계정 상태, 생성·수정 시각"],
        ["사용자", "user_identities", "user_id/provider", "FK·UNIQUE", "외부 인증 공급자 식별자"],
        ["프로필", "user_profiles", "user_id", "FK·UNIQUE", "목표, 기본 시간, 주간 빈도"],
        ["프로필", "user_available_locations", "user_id/location_code", "복합 UNIQUE", "HOME·GYM·OUTDOOR"],
        ["프로필", "user_consents", "user_id/consent_code", "복합 UNIQUE", "현재 동의 상태"],
        ["프로필", "mutation_idempotency_records", "user/endpoint/key", "복합 UNIQUE", "요청 hash와 최초 응답"],
        ["카탈로그", "catalog_versions", "id/version", "UUID PK·UNIQUE", "상태, 승인, 운영 적격"],
        ["카탈로그", "exercises", "id/catalog_version_id", "PK·FK", "운동 코드, 검수 상태"],
        ["카탈로그", "exercise_goal_tag_links", "exercise_id/goal_code", "복합 UNIQUE", "승인 목표·역할 연결"],
        ["카탈로그", "exercise_prescription_profiles", "exercise/goal/level", "복합 UNIQUE", "세트·반복·휴식 기본값"],
        ["루틴", "routines", "id/user_id/version", "PK·FK·UNIQUE", "활성 기간, 요청 시간, 목표"],
        ["루틴", "routine_days", "routine_id/sequence", "복합 UNIQUE", "ROTATION 순환 순서, setup"],
        ["루틴", "routine_items", "routine_day_id/sequence", "복합 UNIQUE", "phase·tier·시간 구성"],
        ["체크인", "daily_contexts", "id/user/date/version", "PK·복합 UNIQUE", "정규화 condition snapshot"],
        ["체크인", "daily_context_discomforts", "context/body_area", "복합 UNIQUE", "불편 부위·심각도"],
        ["체크인", "daily_context_adverse_reactions", "context/reaction", "복합 UNIQUE", "중대한 이상 반응 코드"],
        ["결정", "decision_policy_versions", "id/version_code", "PK·UNIQUE", "ACTIVE·DEPRECATED"],
        ["결정", "decision_runs", "id/input_hash", "PK·INDEX", "입력 snapshot, 버전 조합, 최종 결과"],
        ["결정", "agent_proposals", "run_id/agent_type", "복합 UNIQUE", "4개 Agent 구조화 proposal"],
        ["결정", "plan_candidates", "run_id/candidate_code", "복합 UNIQUE", "공통 후보, 시간, 선택 여부"],
        ["결정", "plan_items", "candidate_id/sequence", "복합 UNIQUE", "승인 운동과 정확한 duration"],
        ["결정", "safety_reviews", "decision_run_id", "FK·UNIQUE", "Safety 상태, veto, reason code"],
        ["결정", "decision_options", "run_id/option_code", "복합 UNIQUE", "FINAL_ROUTINE 또는 REST"],
    ]
    replace_table_rows(doc.tables[3], entity_rows)

    relationship_rows = [
        ["사용자–프로필", "users", "user_profiles", "1:1"],
        ["사용자–루틴", "users", "routines", "1:N"],
        ["루틴–일자", "routines", "routine_days", "1:N"],
        ["일자–항목", "routine_days", "routine_items", "1:N"],
        ["카탈로그–운동", "catalog_versions", "exercises", "1:N"],
        ["사용자–체크인", "users", "daily_contexts", "1:N"],
        ["체크인–불편", "daily_contexts", "daily_context_discomforts", "1:N"],
        ["사용자–결정", "users", "decision_runs", "1:N"],
        ["결정–Agent 제안", "decision_runs", "agent_proposals", "1:4"],
        ["결정–후보", "decision_runs", "plan_candidates", "1:N"],
        ["후보–항목", "plan_candidates", "plan_items", "1:N"],
        ["결정–안전 검수", "decision_runs", "safety_reviews", "1:1"],
        ["결정–공개 선택", "decision_runs", "decision_options", "1:N"],
    ]
    replace_table_rows(doc.tables[4], relationship_rows)

    erd_heading = next(p for p in doc.paragraphs if p.text.strip() == "2.3 ERD")
    # The template places a manual page break after the ERD heading. Removing
    # only that break keeps the heading and inserted diagram on the same page.
    for page_break in list(
        erd_heading._p.iter("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}br")
    ):
        page_break.getparent().remove(page_break)
    erd_paragraph = insert_paragraph_after(erd_heading)
    erd_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    erd_paragraph.add_run().add_picture(str(ERD_PATH), width=Cm(16.0))
    caption = insert_paragraph_after(erd_paragraph, "그림 1. Wave 6 결정 저장 핵심 관계")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    set_paragraph_text(
        next(p for p in doc.paragraphs if p.text.strip() == "테이블 : users"),
        "테이블 : decision_runs (Wave 6 결정 재현성 기준 레코드)",
    )
    decision_run_rows = [
        ["id", "UUID", "✔", "", "✔", "결정 실행 식별자"],
        ["user_id", "UUID", "", "✔", "✔", "users.id, ON DELETE CASCADE"],
        ["local_date", "DATE", "", "", "✔", "사용자 로컬 기준 결정일"],
        ["daily_context_id", "UUID", "", "✔", "✔", "daily_contexts.id, RESTRICT"],
        ["daily_context_version", "INTEGER", "", "", "✔", "조회 당시 check-in version"],
        ["base_routine_id", "UUID", "", "✔", "✔", "routines.id, RESTRICT"],
        ["input_schema_version", "VARCHAR(64)", "", "", "✔", "snapshot schema 식별자"],
        ["input_snapshot", "JSONB", "", "", "✔", "최소 정규화 입력; 직접 식별자 금지"],
        ["input_hash", "VARCHAR(64)", "", "", "✔", "정렬 JSON의 SHA-256"],
        ["catalog_version_id", "UUID", "", "✔", "✔", "사용한 catalog version"],
        ["policy_version_id", "UUID", "", "✔", "✔", "decision policy version"],
        ["safety_rule_version", "VARCHAR(128)", "", "", "✔", "Safety ruleset version"],
        ["duration_rule_version", "VARCHAR(128)", "", "", "✔", "duration rule version"],
        ["graph_version", "VARCHAR(128)", "", "", "✔", "Agent graph version"],
        ["coordinator_version", "VARCHAR(128)", "", "", "✔", "Coordinator version"],
        ["status_code", "VARCHAR(16)", "", "", "✔", "RUNNING·COMPLETED·FAILED·NEEDS_INPUT"],
        ["safety_status_code", "VARCHAR(16)", "", "", "✔", "Safety 결과 원문 보존"],
        ["recommended_action_code", "VARCHAR(32)", "", "", "", "KEEP·DOWNSHIFT·REST 등"],
        ["coordinator_result", "JSONB", "", "", "✔", "구조화 최종 결과 snapshot"],
        ["failure_code", "VARCHAR(128)", "", "", "", "실패 시 machine-readable code"],
        ["created_at/completed_at", "TIMESTAMPTZ", "", "", "", "생성·완료 시각"],
    ]
    replace_table_rows(doc.tables[5], decision_run_rows)

    constraint_rows = [
        ["DB 레벨", "PK", "모든 핵심 테이블", "UUID로 레코드 식별"],
        ["DB 레벨", "FK", "routine·context·decision 계층", "부모 삭제 정책을 CASCADE/RESTRICT로 명시"],
        ["DB 레벨", "UNIQUE", "agent_proposals(run,type)", "Agent 유형 중복 저장 차단"],
        ["DB 레벨", "CHECK", "agent_type/status/action/option", "승인된 안정 코드만 저장"],
        ["DB 레벨", "INDEX", "decision_runs(user,date,hash)", "사용자별 조회와 재현성 검사 지원"],
        ["앱 레벨", "정규 hash", "input_snapshot", "JSON key·proposal 입력 순서와 무관한 SHA-256"],
        ["앱 레벨", "정확한 시간", "선택 plan_candidate", "estimated_seconds = requested_minutes × 60"],
        ["앱 레벨", "Safety veto", "safety_review/final option", "BLOCKED·veto를 성공 routine으로 변형 금지"],
        ["앱 레벨", "필수 proposal", "decision_run", "TRAINING·RECOVERY·SAFETY·FEASIBILITY 정확히 4개"],
        ["앱 레벨", "멱등성", "mutation_idempotency_records", "동일 key+payload는 최초 응답, 불일치는 409"],
        ["저장소", "트랜잭션", "결정 aggregate 전체", "run·proposal·candidate·safety·option 원자 저장"],
        ["저장소", "Fail closed", "DB 저장 실패", "rollback 후 성공 응답을 반환하지 않음"],
        ["보안", "최소 수집", "snapshot·proposal·로그", "이메일·이름·생년월일·토큰·원시 건강정보 금지"],
    ]
    replace_table_rows(doc.tables[6], constraint_rows)

    change_rows = [
        ["2026.08.14", "30기 5팀", "Wave 6 병합 기준 데이터 모델 현행화", "전체", "origin/develop 확인"],
        ["2026.08.14", "채동현", "재현성·4개 proposal·Safety veto 저장 계약 반영", "결정 저장", "문서 작성"],
    ]
    replace_table_rows(doc.tables[7], change_rows)

    doc.core_properties.title = "데이터베이스/저장소 설계 문서 - 30기 5팀"
    doc.core_properties.author = "채동현"
    doc.core_properties.subject = "Wave 6 병합 기준 PostgreSQL 및 저장소 설계"
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
